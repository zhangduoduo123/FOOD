import os
import pandas as pd
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain_community.embeddings import OllamaEmbeddings
from langchain.schema import Document
from docx import Document as DocxDocument
import shutil

# 1. 读取 index.xlsx 获取文件名列表
index_path = "knowledgeApp\doc_process\index.xlsx"
pdf_dir = "pdfs"
df = pd.read_excel(index_path)
file_names = df['书名'].tolist()  
scores = df['权重'].tolist()   # 权重列
classes = df['分类'].tolist()  # 分类列

chroma_db_dir = "chroma_db"
if os.path.exists(chroma_db_dir):
    shutil.rmtree(chroma_db_dir)
    print("已清空 chroma_db 目录，准备重新入库。")

# 2. 遍历文件名，处理每个PDF
all_chunks = []
success_count = 0
matched_files = set()
for file_name, score, cls in zip(file_names, scores, classes):
    # 模糊匹配pdfs目录下的文件（忽略扩展名，包含关系）
    base_name = os.path.splitext(file_name)[0]
    candidates = [f for f in os.listdir(pdf_dir) if base_name in os.path.splitext(f)[0]]
    if not candidates:
        continue
    file_real_name = candidates[0]
    matched_files.add(file_real_name)
    file_ext = os.path.splitext(file_real_name)[-1].lower()
    file_path = os.path.join(pdf_dir, file_real_name)
    try:
        if file_ext == '.pdf':
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_ext == '.docx':
            docx = DocxDocument(file_path)
            text = '\n'.join([para.text for para in docx.paragraphs if para.text.strip()])
            documents = [Document(page_content=text, metadata={})] if text else []
        else:
            print(f"不支持的文件类型，跳过: {file_real_name}")
            continue
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(documents)
    except Exception as e:
        print(f"处理失败，跳过: {file_real_name}，原因: {e}")
        continue
    print(f"成功匹配并入库: {file_real_name} | score: {score} | class: {cls}")
    success_count += 1
    # 4. 每个chunk加score和class字段
    for chunk in chunks:
        if not hasattr(chunk, 'metadata') or chunk.metadata is None:
            chunk.metadata = {}
        chunk.metadata['score'] = score
        chunk.metadata['class'] = cls
    all_chunks.extend(chunks)

# 3. 处理未匹配到index的pdfs目录下的PDF/Word文档
for file_real_name in os.listdir(pdf_dir):
    if file_real_name in matched_files:
        continue
    file_ext = os.path.splitext(file_real_name)[-1].lower()
    if file_ext not in ['.pdf', '.docx']:
        continue
    file_path = os.path.join(pdf_dir, file_real_name)
    try:
        if file_ext == '.pdf':
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_ext == '.docx':
            docx = DocxDocument(file_path)
            text = '\n'.join([para.text for para in docx.paragraphs if para.text.strip()])
            documents = [Document(page_content=text, metadata={})] if text else []
        else:
            continue
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(documents)
    except Exception as e:
        print(f"处理失败，跳过: {file_real_name}，原因: {e}")
        continue
    print(f"未在index中匹配，入库: {file_real_name} | score: 0 | class: 其他")
    # 4. 每个chunk加score=0和class=其他
    for chunk in chunks:
        if not hasattr(chunk, 'metadata') or chunk.metadata is None:
            chunk.metadata = {}
        chunk.metadata['score'] = 0
        chunk.metadata['class'] = '其他'
    all_chunks.extend(chunks)

# 5. 入库（如Chroma）
if all_chunks:
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    vectordb = Chroma.from_documents(
        all_chunks,
        embeddings,
        persist_directory="chroma_db"
    )
    print("入库完成。")
else:
    print("没有可入库的文档。")

print(f"成功入库的文档总数: {success_count}")