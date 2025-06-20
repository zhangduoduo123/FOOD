# -*- coding: utf-8 -*-
from typing import Union, List, Optional, Dict
import os
import json
from pathlib import Path
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain.document_loaders import PyPDFLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.schema import Document
import numpy as np
import pandas as pd
from bs4 import BeautifulSoup
import traceback
from docx import Document as DocxDocument
import time
from functools import lru_cache

# 全局RAG实例
_rag_instance = None

class Config:
    """配置管理类"""
    def __init__(self, config_path: str = "config.json"):
        self.config_path = config_path
        self.default_config = {
            "embedding_model": "nomic-embed-text",
            "llm_model": "deepseek-r1:8b",
            "llm_temperature": 0.7,
            "chunk_size": 1000,
            "chunk_overlap": 200,
            "vector_db_dir": "chroma_db",
            "search_k": 10
        }
        self.config = self._load_config()

    def _load_config(self) -> dict:
        """加载配置文件，如果不存在则创建默认配置"""
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                return self.default_config
        else:
            try:
                with open(self.config_path, 'w', encoding='utf-8') as f:
                    json.dump(self.default_config, f, indent=4, ensure_ascii=False)
                return self.default_config
            except Exception as e:
                return self.default_config

    def get(self, key: str, default=None):
        """获取配置项"""
        return self.config.get(key, default)

class RAGSystem:
    """RAG系统主类"""
    def __init__(self, config_path: str = "config.json"):
        self.config = Config(config_path)
        self.qa_chain = None
        self.conversation_history = []  # 存储最近三次问答
        self.max_history = 3  # 最大历史记录数

    def _check_pdf_path(self, pdf_path: Union[str, List[str]]):
        """检查PDF路径是否有效"""
        if isinstance(pdf_path, list):
            for path in pdf_path:
                if not os.path.exists(path):
                    raise FileNotFoundError(f"文件不存在: {path}")
        elif os.path.isdir(pdf_path):
            if not any(f.endswith(('.pdf', '.csv', '.xls', '.xlsx', '.html', '.docx')) for f in os.listdir(pdf_path)):
                raise ValueError(f"目录中没有可用的知识文件: {pdf_path}")
        elif not os.path.exists(pdf_path):
            raise FileNotFoundError(f"文件不存在: {pdf_path}")

    def _ensure_vector_db_dir(self):
        """确保向量数据库目录存在且有写入权限"""
        db_dir = self.config.get("vector_db_dir")
        try:
            Path(db_dir).mkdir(parents=True, exist_ok=True)
            test_file = Path(db_dir) / ".test"
            test_file.touch()
            test_file.unlink()
        except Exception as e:
            raise RuntimeError(f"向量数据库目录权限错误: {str(e)}")

    def _load_documents(self, pdf_path: Union[str, List[str]]) -> List[Document]:
        """加载文档文件"""
        documents = []
        
        if isinstance(pdf_path, list):
            for file_path in pdf_path:
                if os.path.exists(file_path):
                    docs = self._load_single_document(file_path)
                    documents.extend(docs)
        elif os.path.isdir(pdf_path):
            for file_name in os.listdir(pdf_path):
                file_path = os.path.join(pdf_path, file_name)
                if os.path.isfile(file_path):
                    docs = self._load_single_document(file_path)
                    documents.extend(docs)
        elif os.path.isfile(pdf_path):
            docs = self._load_single_document(pdf_path)
            documents.extend(docs)
        
        return documents

    def _load_single_document(self, file_path: str) -> List[Document]:
        """加载单个文档文件"""
        file_ext = os.path.splitext(file_path)[-1].lower()
        file_name = os.path.basename(file_path)
        
        try:
            if file_ext == '.pdf':
                loader = PyPDFLoader(file_path)
                documents = loader.load()
                # 添加文档名称到元数据
                for doc in documents:
                    if not hasattr(doc, 'metadata') or doc.metadata is None:
                        doc.metadata = {}
                    doc.metadata['name'] = file_name
                return documents
            elif file_ext == '.docx':
                docx = DocxDocument(file_path)
                text = '\n'.join([para.text for para in docx.paragraphs if para.text.strip()])
                if text:
                    doc = Document(page_content=text, metadata={'name': file_name})
                    return [doc]
                return []
            elif file_ext in ['.csv', '.xls', '.xlsx']:
                if file_ext == '.csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                text_content = df.to_string(index=False)
                doc = Document(page_content=text_content, metadata={'name': file_name})
                return [doc]
            elif file_ext == '.html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                text_content = soup.get_text()
                doc = Document(page_content=text_content, metadata={'name': file_name})
                return [doc]
            else:
                print(f"不支持的文件类型: {file_ext}")
                return []
        except Exception as e:
            print(f"加载文档失败 {file_path}: {str(e)}")
            return []

    def setup_qa_chain(self, pdf_path: Union[str, List[str]] = "pdfs"):
        """设置QA链，支持处理PDF、表格文件、文件列表或目录"""
        try:
            self._check_pdf_path(pdf_path)
            self._ensure_vector_db_dir()

            db_dir = self.config.get("vector_db_dir")
            embeddings = OllamaEmbeddings(model=self.config.get("embedding_model"))
            vectordb = None

            # 判断ChromaDB是否已存在
            if os.path.exists(db_dir) and os.listdir(db_dir):
                vectordb = Chroma(persist_directory=db_dir, embedding_function=embeddings)
            else:
                documents = self._load_documents(pdf_path)
                # 统一转换为Document对象
                documents = [
                    doc if isinstance(doc, Document) else Document(page_content=doc.get("page_content", ""), metadata=doc.get("metadata", {}))
                    for doc in documents
                ]
                
                # 文本分割
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.config.get("chunk_size"),
                    chunk_overlap=self.config.get("chunk_overlap")
                )
                texts = text_splitter.split_documents(documents)
                
                # 初始化ChromaDB
                vectordb = Chroma.from_documents(
                    texts,
                    embeddings,
                    persist_directory=db_dir
                )

            # 初始化生成模型
            llm = ChatOllama(
                model=self.config.get("llm_model"),
                temperature=self.config.get("llm_temperature")
            )

            # 定义自定义的提示模板
            qa_prompt = PromptTemplate(
                template="你是一个专业的知识助手，根据以下上下文回答用户的问题。如果不知道答案，请诚实说明。\n"
                         "上下文信息：\n{context}\n\n"
                         "用户问题：{question}\n\n"
                         "请根据上述上下文，提供准确、全面的回答。"
                         "如果不知道答案，请诚实说明。回答要结构清晰，重点突出。\n\n",
                input_variables=["context", "question"]
            )

            self.qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=vectordb.as_retriever(
                    search_kwargs={"k": self.config.get("search_k")}
                ),
                return_source_documents=True,
                chain_type_kwargs={"prompt": qa_prompt}
            )
            return True
        except Exception as e:
            print("[setup_qa_chain] 异常：", e)
            return False

    def query_documents(self, query: str, pdf_path: Optional[Union[str, List[str]]] = None) -> Dict:
        """查询文档并返回答案"""
        try:
            start_time = time.time()
            
            if pdf_path is not None or self.qa_chain is None:
                success = self.setup_qa_chain(pdf_path or "pdfs")
                if not success:
                    return {"status": "error", "message": "Failed to setup QA chain"}
            
            if self.qa_chain is None:
                return {"status": "error", "message": "QA chain not initialized"}
            
            # 召回文档
            retriever = self.qa_chain.retriever
            docs = retriever.get_relevant_documents(query)
            
            # 打印top-k结果
            print(f"\n=== 查询: {query} ===")
            print(f"检索到 {len(docs)} 个相关文档片段:")
            print("-" * 50)
            
            for i, doc in enumerate(docs):
                metadata = doc.metadata
                doc_name = metadata.get('name', '未知文档')
                print(f"[{i+1}] 文档: {doc_name}")
                print(f"完整内容: {doc.page_content}")
                print("-----------------------------")
            
            # 让大模型判断文档是否相关
            if docs:
                # 构建文档内容用于判断
                docs_content = "\n\n".join([f"文档{i+1}: {doc.page_content}" for i, doc in enumerate(docs)])
                
                # 让大模型判断相关性
                relevance_prompt = f"""请判断以下检索到的文档是否与用户问题相关。

                        用户问题：{query}

                        检索到的文档内容：
                        {docs_content}

                        请回答：
                        1. 这些文档是否与用户问题相关？（是/否）
                        2. 如果相关，请说明原因；如果不相关，请说明原因。

                        回答格式：
                        相关性：是/否
                        原因：[详细说明]"""

                relevance_judgment = self.qa_chain.combine_documents_chain.llm_chain.llm.predict(relevance_prompt)
                print(f"\n=== 大模型相关性判断 ===")
                print(relevance_judgment)
                print("=" * 50)
                
                # 解析大模型的判断结果
                print(f"DEBUG: 原始判断结果: {relevance_judgment}")
                
                # 更精确的判断逻辑
                is_relevant = False
                if "相关性：是" in relevance_judgment:
                    is_relevant = True
                elif "相关性：否" in relevance_judgment:
                    is_relevant = False
                else:
                    # 如果没有明确的格式，尝试其他判断方式
                    if "是" in relevance_judgment[:20] and "否" not in relevance_judgment[:20]:
                        is_relevant = True
                    elif "否" in relevance_judgment[:20]:
                        is_relevant = False
                    else:
                        # 默认认为不相关
                        is_relevant = False
                
                
                
                if is_relevant:
                    # 文档相关，使用RAG方式回答
                    print("大模型判断：文档相关，使用RAG方式回答")
                    
                    # 获取历史对话上下文
                    history_context = self.get_history_context()
                    
                    # 构建包含历史对话的提示
                    if history_context:
                        print(f"\n=== 历史对话上下文 ===")
                        print(history_context)
                        print("=" * 50)
                        
                        # 手动构建包含历史的提示
                        docs_content = "\n\n".join([doc.page_content for doc in docs])
                        full_prompt = f"""你是一个专业的知识助手，根据以下上下文回答用户的问题。如果不知道答案，请诚实说明。

                                    历史对话上下文：
                                    {history_context}

                                    当前检索到的文档信息：
                                    {docs_content}

                                    用户问题：{query}

                                    请根据上述上下文和历史对话，提供准确、全面的回答。如果不知道答案，请诚实说明。回答要结构清晰，重点突出。

                                    回答："""
                        
                        # 直接使用LLM生成回答
                        result = self.qa_chain.combine_documents_chain.llm_chain.llm.predict(full_prompt)
                    else:
                        # 没有历史对话，使用标准RAG方式
                        result = self.qa_chain.combine_documents_chain.run({"input_documents": docs, "question": query})
                    
                    # 收集引用的文档名称
                    cited_documents = set()
                    for doc in docs:
                        metadata = doc.metadata
                        doc_name = metadata.get('name', '未知文档')
                        cited_documents.add(doc_name)
                    
                    # 构建引用字符串
                    citation_text = "\n\n数据引用于：" + "，".join(cited_documents)
                    result = result + citation_text
                else:
                    # 文档不相关，直接使用大模型回答
                    print("大模型判断：文档不相关，直接使用大模型回答")
                    
                    # 获取历史对话上下文
                    history_context = self.get_history_context()
                    
                    if history_context:
                        # 包含历史对话的直接回答
                        full_prompt = f"""历史对话上下文：
                                    {history_context}

                                    当前问题：{query}

                                    请根据历史对话上下文回答当前问题。如果不知道答案，请诚实说明。

                                    回答："""
                        result = self.qa_chain.combine_documents_chain.llm_chain.llm.predict(full_prompt)
                    else:
                        # 没有历史对话，直接回答
                        result = self.qa_chain.combine_documents_chain.llm_chain.llm.predict(query)
            
            # 将当前问答对添加到历史记录
            self.add_to_history(query, result)
            
            end_time = time.time()
            total_time = end_time - start_time
            
            # 提取文档来源信息
            sources_info = []
            if docs:  # 只有在有文档时才处理来源信息
                for i, doc in enumerate(docs):
                    metadata = doc.metadata
                    source_info = {
                        "index": i + 1,
                        "document_name": metadata.get('name', '未知文档'),
                        "content_preview": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content
                    }
                    sources_info.append(source_info)
                
                # 统计文档来源分布
                document_distribution = {}
                for doc in docs:
                    doc_name = doc.metadata.get('name', '未知文档')
                    document_distribution[doc_name] = document_distribution.get(doc_name, 0) + 1
            else:
                # 没有文档时，设置空的来源信息
                document_distribution = {}
            
            return {
                "status": "success",
                "answer": result,
                "processing_time": total_time,
                "sources": [doc.page_content for doc in docs] if docs else [],
                "topk_contents": [doc.page_content for doc in docs] if docs else [],
                "sources_info": sources_info,
                "document_distribution": document_distribution,
                "total_sources": len(docs)
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}

    def add_to_history(self, question: str, answer: str):
        """添加问答对到历史记录"""
        qa_pair = {
            "question": question,
            "answer": answer,
            "timestamp": time.time()
        }
        
        self.conversation_history.append(qa_pair)
        
        # 保持最近三次记录
        if len(self.conversation_history) > self.max_history:
            self.conversation_history = self.conversation_history[-self.max_history:]
        
        print(f"已添加问答到历史，当前历史记录数: {len(self.conversation_history)}")
    
    def get_history_context(self) -> str:
        """获取历史对话上下文"""
        if not self.conversation_history:
            return ""
        
        context_parts = []
        for i, qa in enumerate(self.conversation_history, 1):
            context_parts.append(f"第{i}次对话:")
            context_parts.append(f"问题: {qa['question']}")
            context_parts.append(f"回答: {qa['answer']}")
            context_parts.append("")
        
        return "\n".join(context_parts)
    
    def clear_history(self):
        """清空历史记录"""
        self.conversation_history = []
        print("历史记录已清空")

def setup_qa_chain(pdf_path: Union[str, List[str]] = "pdfs") -> Dict:
    """设置QA链，支持处理单个文件、文件列表或目录"""
    global _rag_instance
    try:
        if _rag_instance is None:
            _rag_instance = RAGSystem()
        
        success = _rag_instance.setup_qa_chain(pdf_path)
        if success:
            return {"status": "success", "message": "QA chain setup completed"}
        else:
            return {"status": "error", "message": "Failed to setup QA chain"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

def query_documents(query: str, pdf_path: Optional[Union[str, List[str]]] = None) -> Dict:
    """查询文档并返回答案"""
    global _rag_instance
    try:
        if _rag_instance is None:
            _rag_instance = RAGSystem()
        
        return _rag_instance.query_documents(query, pdf_path)
    except Exception as e:
        return {"status": "error", "message": str(e)}

def clear_conversation_history() -> Dict:
    """清空对话历史"""
    global _rag_instance
    try:
        if _rag_instance is None:
            _rag_instance = RAGSystem()
        
        _rag_instance.clear_history()
        return {"status": "success", "message": "Conversation history cleared"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

def get_conversation_history() -> Dict:
    """获取当前对话历史"""
    global _rag_instance
    try:
        if _rag_instance is None:
            _rag_instance = RAGSystem()
        
        history = _rag_instance.conversation_history
        return {
            "status": "success", 
            "history": history,
            "count": len(history)
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    # 示例调用
    result = query_documents("什么是人工智能？", "pdfs")
    print(result) 